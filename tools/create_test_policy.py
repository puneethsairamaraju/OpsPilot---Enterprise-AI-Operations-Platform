from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path("sample_uploads/Acme_Hybrid_Work_and_Equipment_Policy.docx")
BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 103, 110)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    paragraph.add_run(text)


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "ACME OPERATIONS  |  PEOPLE POLICY"
    header.style = styles["Normal"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.text = "Internal use  |  Policy HR-204  |  Effective July 1, 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("INTERNAL POLICY")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("Hybrid Work and\nEquipment Policy")
    run.bold = True
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(
        "Eligibility, scheduling, home-office support, security, and approval requirements"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    metadata = document.add_table(rows=4, cols=2)
    metadata.autofit = False
    metadata.columns[0].width = Inches(1.35)
    metadata.columns[1].width = Inches(5.15)
    values = [
        ("Policy owner", "People Operations"),
        ("Effective date", "July 1, 2026"),
        ("Review cycle", "Annual"),
        ("Applies to", "Full-time employees in the United States"),
    ]
    for row, (label, value) in zip(metadata.rows, values, strict=True):
        row.cells[0].width = Inches(1.35)
        row.cells[1].width = Inches(5.15)
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    document.add_heading("1. Purpose", level=1)
    document.add_paragraph(
        "This policy establishes a consistent hybrid-work model while preserving "
        "customer coverage, team collaboration, information security, and employee wellbeing."
    )

    document.add_heading("2. Standard hybrid schedule", level=1)
    document.add_paragraph(
        "Eligible employees may work remotely up to three days per week. Tuesday and "
        "Wednesday are company-wide collaboration days and must normally be worked from "
        "the employee's assigned office."
    )
    add_bullet(
        document,
        "Managers may designate one additional in-office day when customer or project needs require it.",
    )
    add_bullet(
        document,
        "A permanent schedule exception requires approval from the department director and People Operations.",
    )
    add_bullet(
        document,
        "Employees must remain available during their team's published core hours of 10:00 a.m. to 3:00 p.m. local time.",
    )

    document.add_heading("3. International remote work", level=1)
    document.add_paragraph(
        "Working from another country is limited to 15 calendar days per year. Employees "
        "must submit a request at least 30 days before travel. Approval is required from "
        "the employee's department director, People Operations, and Information Security."
    )

    document.add_heading("4. Home-office equipment", level=1)
    document.add_paragraph(
        "The company provides a laptop and an annual home-office allowance of 800 dollars. "
        "The allowance may be used for a monitor, keyboard, mouse, ergonomic chair, desk, "
        "headset, or task lighting."
    )
    add_bullet(
        document,
        "Individual purchases above 300 dollars require manager approval before purchase.",
    )
    add_bullet(
        document,
        "Receipts must be submitted through the expense system within 10 calendar days.",
    )
    add_bullet(
        document,
        "Internet service, mobile-phone plans, decorations, and household utilities are not reimbursable.",
    )

    document.add_heading("5. Security requirements", level=1)
    document.add_paragraph(
        "Employees must use the company VPN when accessing internal systems outside an "
        "office. Company information may not be stored on personal devices or printed in "
        "public locations. A lost company device must be reported to Security Operations "
        "within one hour of discovery."
    )

    document.add_heading("6. Questions and exceptions", level=1)
    document.add_paragraph(
        "Questions should be sent to People Operations. Exceptions are reviewed case by "
        "case and do not establish a permanent precedent for other employees."
    )

    document.core_properties.title = "Acme Hybrid Work and Equipment Policy"
    document.core_properties.subject = "Sample knowledge-base upload document"
    document.core_properties.author = "Acme Operations"
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
